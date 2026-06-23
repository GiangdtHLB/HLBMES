#!/usr/bin/env python3
"""Chuyển tài liệu Markdown của MES sang .docx có định dạng chuyên nghiệp.

Hỗ trợ: heading (#/##/###), đoạn văn, blockquote (>), bảng (| |), danh sách
gạch đầu dòng (-) và đánh số (1.), code fence (```), inline **bold** / `code` /
[label](href). Có trang bìa + mục lục (TOC field) + header/footer số trang.

Dùng: python md2docx.py input.md output.docx "Tiêu đề bìa" "Phụ đề"
"""
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---- Bảng màu thương hiệu ----
NAVY = RGBColor(0x1E, 0x27, 0x61)
TEAL = RGBColor(0x02, 0x80, 0x90)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0xF0, 0xF4, 0xF8)
HDR_FILL = "1E2761"
ALT_FILL = "EEF2F8"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def set_cell_bg(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def add_table_borders(table, color="BBBBBB", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, base_size=None, base_color=None, base_bold=False):
    """Phân tích **bold**, `code`, [label](href) và thêm run vào paragraph."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("]")]
            r = paragraph.add_run(label)
            r.font.color.rgb = TEAL
            r.underline = True
        else:
            r = paragraph.add_run(part)
        if base_size:
            r.font.size = base_size
        if base_color is not None:
            r.font.color.rgb = base_color
        if base_bold:
            r.bold = True


def strip_md(text):
    """Bỏ ký hiệu markdown để lấy text thuần (dùng trong ô bảng đơn giản)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Nhấn Ctrl+A rồi F9 trong Word để cập nhật mục lục."
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run("Trang ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run2 = p.add_run(); run2.font.size = Pt(8); run2.font.color.rgb = GREY
    run2._r.append(fld1); run2._r.append(instr); run2._r.append(fld2)


def style_document(document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 13, NAVY),
                              ("Heading 3", 11.5, TEAL)):
        st = document.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10 if size >= 13 else 6)
        st.paragraph_format.space_after = Pt(4)


def build_cover(document, title, subtitle):
    for _ in range(3):
        document.add_paragraph()
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🍺 MES Bia Hạ Long"); r.font.size = Pt(20); r.bold = True; r.font.color.rgb = TEAL
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhà Máy Đông Mai"); r.font.size = Pt(12); r.font.color.rgb = GREY
    for _ in range(2):
        document.add_paragraph()
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY
    if subtitle:
        p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle); r.font.size = Pt(13); r.font.color.rgb = GREY
    for _ in range(6):
        document.add_paragraph()
    for line in ("Tài liệu hướng dẫn nội bộ", "Phiên bản phần mềm: 0.1.0-mvp",
                 "Ngày phát hành: 23/06/2026"):
        p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(10.5); r.font.color.rgb = GREY


def parse_table(lines, i, document):
    """Đọc khối bảng markdown bắt đầu tại dòng i; trả (table_built, next_i)."""
    rows = []
    j = i
    while j < len(lines) and lines[j].strip().startswith("|"):
        rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
        j += 1
    # rows[1] là dòng phân cách ---
    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    ncol = len(header)
    table = document.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    hdr = table.rows[0].cells
    for c, txt in enumerate(header):
        set_cell_bg(hdr[c], HDR_FILL)
        set_cell_margins(hdr[c])
        para = hdr[c].paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(strip_md(txt))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for ri, row in enumerate(body):
        cells = table.add_row().cells
        for c in range(ncol):
            val = row[c] if c < len(row) else ""
            if ri % 2 == 1:
                set_cell_bg(cells[c], ALT_FILL)
            set_cell_margins(cells[c])
            para = cells[c].paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_inline(para, val)
            for rn in para.runs:
                if rn.font.size is None:
                    rn.font.size = Pt(9)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return j


def convert(md_path, out_path, title, subtitle):
    with open(md_path, encoding="utf-8") as f:
        raw = f.read().split("\n")

    document = Document()
    for section in document.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    style_document(document)

    build_cover(document, title, subtitle)
    document.add_page_break()
    h = document.add_heading("Mục lục", level=1)
    add_toc(document)
    document.add_page_break()
    add_page_number_footer(document.sections[0])

    # Bỏ qua H1 đầu tiên (đã dùng làm bìa)
    skipped_first_h1 = False
    i = 0
    n = len(raw)
    in_code = False
    code_buf = []
    while i < n:
        line = raw[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                # đóng code block
                cp = document.add_paragraph()
                cp.paragraph_format.left_indent = Pt(8)
                for ci, cl in enumerate(code_buf):
                    if ci:
                        cp.add_run().add_break()
                    rr = cp.add_run(cl)
                    rr.font.name = MONO_FONT
                    rr.font.size = Pt(8.5)
                    rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("# "):
            if not skipped_first_h1:
                skipped_first_h1 = True
                i += 1
                continue
            document.add_heading(strip_md(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            document.add_heading(strip_md(stripped[3:]), level=2)
        elif stripped.startswith("### "):
            document.add_heading(strip_md(stripped[4:]), level=3)
        elif stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|?$", raw[i + 1].strip()):
            i = parse_table(raw, i, document)
            continue
        elif stripped.startswith("> "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_before = Pt(2)
            add_inline(p, stripped[2:], base_color=GREY)
            for rn in p.runs:
                rn.italic = True
        elif re.match(r"^[-*] ", stripped):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = document.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped == "---":
            pass  # bỏ qua đường kẻ ngang
        elif stripped == "":
            pass
        else:
            p = document.add_paragraph()
            add_inline(p, stripped)
        i += 1

    document.save(out_path)
    print(f"✓ {out_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4] if len(sys.argv) > 4 else "")
