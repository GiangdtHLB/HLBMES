"""Convert docs/03-HUONG-DAN-SU-DUNG.md -> docs/03-HUONG-DAN-SU-DUNG.docx
Styling matched by hand against docs/01-KIEN-TRUC-CHUAN.docx (colors, fonts, table borders).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_SECTION

NAVY = RGBColor(0x1E, 0x27, 0x61)
TEAL = RGBColor(0x02, 0x80, 0x90)
GRAY = RGBColor(0x55, 0x55, 0x55)
BORDER = "BBBBBB"
ZEBRA = "EEF2F8"
HEADER_FILL = "1E2761"

SRC = Path("docs/03-HUONG-DAN-SU-DUNG.md")
OUT = Path("docs/03-HUONG-DAN-SU-DUNG.docx")


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER)
        borders.append(el)
    tcPr.append(borders)


def set_cell_shading(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, start=110, end=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def shade_paragraph(paragraph, hexcolor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def add_page_number_field(paragraph, prefix="Trang "):
    run = paragraph.add_run(prefix)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\[[^\]]*?\]\([^)]*?\))")


def add_inline_runs(paragraph, text, base_size=10.5, base_color=None, base_bold=False, base_italic=False):
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]*)\]\(([^)]*)\)", part)
            label = m.group(1) if m else part
            run = paragraph.add_run(label)
        else:
            run = paragraph.add_run(part)
        if base_size:
            run.font.size = Pt(base_size)
        if base_color and not (part.startswith("`") and part.endswith("`")):
            run.font.color.rgb = base_color
        if base_bold:
            run.bold = True
        if base_italic:
            run.italic = True


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = TEAL
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(10.5)
    ln = doc.styles["List Number"]
    ln.font.name = "Calibri"
    ln.font.size = Pt(10.5)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)


def build_cover(doc, title, subtitle, meta_lines):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🍺 MES Bia Hạ Long")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhà máy Đông Mai")
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = TEAL
    r.bold = True

    for _ in range(4):
        doc.add_paragraph()

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, line, base_size=9.5, base_color=GRAY)

    doc.add_page_break()


def parse_pipe_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells = []
    buf = ""
    esc = False
    for ch in line:
        if esc:
            buf += ch
            esc = False
        elif ch == "\\":
            esc = True
        elif ch == "|":
            cells.append(buf.strip())
            buf = ""
        else:
            buf += ch
    cells.append(buf.strip())
    return cells


def is_sep_row(cells):
    return all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())


def add_table(doc, rows):
    if not rows:
        return
    header = rows[0]
    body = rows[1:]
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        cell = hdr_cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        add_inline_runs(para, text, base_size=9.5, base_color=RGBColor(0xFF, 0xFF, 0xFF), base_bold=True)

    for ri, row in enumerate(body):
        cells_row = table.add_row().cells
        for i in range(ncols):
            text = row[i] if i < len(row) else ""
            cell = cells_row[i]
            set_cell_border(cell)
            set_cell_margins(cell)
            if ri % 2 == 1:
                set_cell_shading(cell, ZEBRA)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_inline_runs(para, text, base_size=9)


def add_code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        shade_paragraph(p, "F2F2F2")
        r = p.add_run(ln if ln.strip() else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER)
    pbdr.append(bottom)
    pPr.append(pbdr)


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    setup_styles(doc)

    # ---- parse front matter (title + intro blockquotes before first "---") ----
    i = 0
    title_line = lines[0].lstrip("# ").strip()
    i = 1
    subtitle = ""
    meta = []
    while i < len(lines) and lines[i].strip() != "---":
        raw = lines[i].strip()
        if raw.startswith("**") and raw.endswith("**"):
            subtitle = raw.strip("*")
        elif raw.startswith(">"):
            meta.append(raw.lstrip(">").strip())
        i += 1
    # skip the "---" itself
    if i < len(lines) and lines[i].strip() == "---":
        i += 1

    build_cover(doc, title_line, subtitle, [m for m in meta if m])

    first_h2 = True
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, stripped[4:].strip(), base_size=11.5, base_color=TEAL, base_bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            if not first_h2:
                doc.add_page_break()
            first_h2 = False
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, stripped[3:].strip(), base_size=13, base_color=NAVY, base_bold=True)
            i += 1
            continue

        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, block)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [parse_pipe_row(l) for l in table_lines]
            if len(rows) >= 2 and is_sep_row(rows[1]):
                rows = [rows[0]] + rows[2:]
            add_table(doc, rows)
            continue

        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            joined = " ".join(b for b in block if b)
            add_inline_runs(p, joined, base_size=9.5, base_color=GRAY, base_italic=True)
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
